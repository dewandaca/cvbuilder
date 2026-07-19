# -*- coding: utf-8 -*-
"""Generate 'BAB 3 - Metode Penelitian (revisi 2).docx' from the revisi version,
inserting code blocks (Kode Program) and expanded Prompt Engineering subsections."""
import copy
import shutil
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'BAB 3 - Metode Penelitian (revisi).docx'
DST = 'BAB 3 - Metode Penelitian (revisi 2).docx'

shutil.copyfile(SRC, DST)
doc = Document(DST)

FIRST_LINE_INDENT = Emu(450215)
SPACE_AFTER = Emu(76200)

# ---------------------------------------------------------------- helpers

def _fmt_body(p, indent=True, justify=True):
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = SPACE_AFTER
    if justify:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        pf.first_line_indent = FIRST_LINE_INDENT


def add_runs_with_italic(p, text):
    """Text uses *...* markers for italic runs."""
    parts = text.split('*')
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        if i % 2 == 1:
            r.italic = True


def insert_before(anchor):
    return anchor.insert_paragraph_before()


def body_par(anchor, text, indent=True):
    p = insert_before(anchor)
    _fmt_body(p, indent=indent)
    add_runs_with_italic(p, text)
    return p


def heading_par(anchor, text):
    p = insert_before(anchor)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = SPACE_AFTER
    r = p.add_run(text)
    r.bold = True
    return p


def caption_par(anchor, text):
    p = insert_before(anchor)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Emu(38100)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def _shade_and_border(p):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    pbdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '808080')
        pbdr.append(el)
    pPr.append(pbdr)


def code_par(anchor, code_text):
    p = insert_before(anchor)
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(4)
    pf.space_after = Pt(10)
    pf.left_indent = Pt(6)
    pf.right_indent = Pt(6)
    _shade_and_border(p)
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), 'Consolas')
        rfonts.set(qn('w:hAnsi'), 'Consolas')
        rfonts.set(qn('w:cs'), 'Consolas')
        if i < len(lines) - 1:
            r.add_break()
    return p

# ---------------------------------------------------------------- anchors

anchor_36 = None
anchor_close = None
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('3.6') and 'Prompt Engineering' in t:
        anchor_36 = p
    if t.startswith('Rancangan prompt tersebut menjadi dasar'):
        anchor_close = p
assert anchor_36 is not None and anchor_close is not None

# ================================================================
# 1) Insert into 3.5 (before heading 3.6): Kode Program 3.1
# ================================================================

body_par(anchor_36,
         'Secara teknis, komunikasi dengan LLM diimplementasikan sepenuhnya pada sisi '
         '*server* melalui *Server Action* pada Next.js sehingga *API key* tidak pernah '
         'terekspos ke sisi klien. Setiap permintaan dikirimkan ke *endpoint chat completion* '
         'Groq API dalam bentuk dua pesan, yaitu *system message* yang berisi instruksi hasil '
         'perancangan *prompt engineering* dan *user message* yang berisi teks asli dari '
         'pengguna. Potongan kode pemanggilan LLM pada sistem ditunjukkan pada '
         'Kode Program 3.1 berikut.')

caption_par(anchor_36, 'Kode Program 3.1: Pemanggilan LLM melalui Groq API (src/app/actions.ts)')

code_par(anchor_36,
"""const groq = new Groq({ apiKey: process.env.GROQ_API_KEY });

const chatCompletion = await groq.chat.completions.create({
  messages: [
    { role: 'system', content: systemPrompt }, // instruksi dari sistem
    { role: 'user', content: text },           // teks asli dari pengguna
  ],
  model: 'openai/gpt-oss-120b',
  temperature: 0.5, // 0.1 dan max_tokens: 4096 pada fitur impor CV
});

const result = chatCompletion.choices[0]?.message?.content || '';""")

body_par(anchor_36,
         'Berdasarkan Kode Program 3.1, perilaku model dikendalikan melalui parameter '
         '*temperature* yang berbeda pada setiap fitur. Fitur peningkatan penulisan '
         'menggunakan nilai 0,5 agar model tetap memiliki keleluasaan dalam memilih variasi '
         'kata namun tidak menyimpang dari instruksi, sedangkan fitur impor CV menggunakan '
         'nilai 0,1 disertai batas keluaran 4.096 token agar hasil ekstraksi bersifat '
         'deterministik dan konsisten mengikuti struktur JSON yang ditetapkan. Pemisahan '
         'antara *system message* dan *user message* juga memastikan bahwa teks pengguna '
         'diperlakukan murni sebagai data yang diolah, bukan sebagai instruksi tambahan '
         'bagi model.')

# ================================================================
# 2) Expand 3.6: subsections before the closing paragraph
# ================================================================

# --- transition note appended to paragraph 93 (the one before anchor_close is
#     "Berdasarkan Tabel 3.4, ..."); add pointer sentence via new paragraph
# (kept simple: sentence lives at start of 3.6.1 intro instead)

# ---------------- 3.6.1 ----------------
heading_par(anchor_close, '3.6.1  Prompt Peningkatan Ringkasan Diri (Bahasa Indonesia)')

body_par(anchor_close,
         'Fitur peningkatan ringkasan diri (*summary*) bertugas menulis ulang deskripsi diri '
         'yang ditulis pengguna dengan bahasa sehari-hari menjadi paragraf ringkasan yang '
         'profesional dalam Bahasa Indonesia baku. *System prompt* yang digunakan pada fitur '
         'ini ditunjukkan pada Kode Program 3.2 berikut.')

caption_par(anchor_close, 'Kode Program 3.2: System Prompt Peningkatan Ringkasan Diri (Bahasa Indonesia)')

code_par(anchor_close,
"""You are a seasoned career coach who writes like a real person, not a
machine. Rewrite the user's professional summary in professional
Indonesian (Bahasa Indonesia baku) so it reads like something a
confident human professional would actually say about themselves.

Rules:
- Write 4-6 well-constructed sentences that flow naturally. Do NOT make
  it too short or overly compressed.
- Use active voice and concrete language. Mention specific domains,
  tools, or achievements the user provided — do not invent new ones.
- Vary your sentence openings: do NOT start every sentence with the same
  structure. Mix simple and compound sentences.
- Avoid overused AI filler words such as "berdedikasi", "berpengalaman
  luas", "passionate", "proven track record", "leveraging",
  "spearheading", "cutting-edge", "innovative solutions". Write the way
  a real Indonesian professional would describe themselves in a formal
  setting.
- Keep the tone confident but grounded — not boastful, not robotic.
- Make it ATS-friendly by naturally incorporating relevant keywords from
  the user's input.
- Return ONLY the rewritten summary paragraph. No headings, no labels,
  no conversational filler.""")

body_par(anchor_close,
         'Berdasarkan Kode Program 3.2, prompt tersebut tersusun atas empat komponen utama. '
         'Pertama, pemberian peran (*role*) sebagai *career coach* berpengalaman agar model '
         'mengadopsi sudut pandang dan gaya bahasa seorang profesional, bukan gaya bahasa '
         'mesin. Kedua, aturan struktur yang mewajibkan keluaran berupa 4 sampai 6 kalimat '
         'dengan kalimat aktif serta variasi pembuka kalimat, sehingga ringkasan tidak '
         'terkesan monoton dan terlalu singkat. Ketiga, batasan anti-halusinasi yang melarang '
         'model menambahkan domain, perangkat, maupun pencapaian yang tidak disebutkan '
         'pengguna, sekaligus daftar larangan kata klise khas keluaran AI seperti '
         '"berdedikasi" dan "*proven track record*" agar hasil terdengar alami. Keempat, '
         'aturan format keluaran yang mewajibkan model mengembalikan hanya paragraf hasil '
         'penulisan ulang tanpa judul, label, maupun kalimat pembuka, sehingga keluaran dapat '
         'langsung dimasukkan ke dalam *form* tanpa penyuntingan tambahan. Aspek keramahan '
         'ATS juga ditegaskan dengan meminta model menyisipkan kata kunci relevan dari '
         'masukan pengguna secara alami.')

# ---------------- 3.6.2 ----------------
heading_par(anchor_close, '3.6.2  Prompt Peningkatan Deskripsi Pengalaman, Proyek, dan Pendidikan (Bahasa Indonesia)')

body_par(anchor_close,
         'Fitur peningkatan deskripsi digunakan pada bagian pengalaman kerja, proyek, '
         'organisasi, dan pendidikan untuk mengubah tulisan pengguna menjadi poin-poin '
         '(*bullet points*) yang diawali kata kerja aksi dan ramah ATS. *System prompt* '
         'yang digunakan ditunjukkan pada Kode Program 3.3 berikut.')

caption_par(anchor_close, 'Kode Program 3.3: System Prompt Peningkatan Deskripsi (Bahasa Indonesia)')

code_par(anchor_close,
"""You are a professional Resume Writer who writes like a real person, not
a machine. Improve the user's work/project/education description into
strong, ATS-friendly bullet points written in professional and formal
Indonesian (Bahasa Indonesia baku).

Rules:
- Start each bullet with a strong action verb.
- NEVER invent, fabricate, or assume numbers, percentages, metrics, or
  statistics that the user did not explicitly provide. If the user wrote
  "meningkatkan penjualan", do NOT add "sebesar 20%" or any made-up
  figure. Only include quantifiable data if the user's original text
  already contains it.
- Keep each bullet descriptive and specific based on what the user
  actually wrote — add professional phrasing, not fictional
  achievements.
- Avoid generic AI filler like "secara signifikan", "secara efektif",
  "berdedikasi". Be concrete and natural.
${bulletRule}
- Return ONLY the bullet text, one bullet per line. Do NOT prefix
  bullets with any marker like '•', '-', or '*'. Do NOT output the
  literal characters '\\n' — just use actual line breaks. No
  conversational filler.""")

body_par(anchor_close,
         'Berbeda dengan prompt ringkasan diri, prompt pada Kode Program 3.3 menitikberatkan '
         'pada dua hal. Pertama, aturan anti-fabrikasi yang sangat tegas: model dilarang '
         'keras mengarang angka, persentase, metrik, atau statistik apa pun yang tidak '
         'ditulis pengguna. Aturan ini penting karena LLM cenderung "mempercantik" deskripsi '
         'pekerjaan dengan angka pencapaian fiktif (misalnya menambahkan "sebesar 20%"), '
         'yang dapat membuat isi CV menjadi tidak jujur. Kedua, aturan format keluaran yang '
         'melarang model menuliskan penanda *bullet* maupun karakter literal "\\n", karena '
         'penataan *bullet* dilakukan oleh sistem pada tahap pasca-pemrosesan.')

body_par(anchor_close,
         'Pada bagian ${bulletRule} terdapat aturan dinamis yang disusun oleh *prompt '
         'generator* berdasarkan jumlah poin yang ditulis pengguna. Sistem terlebih dahulu '
         'menghitung jumlah baris bermakna pada masukan pengguna, kemudian menyisipkan '
         'aturan jumlah poin ke dalam prompt sebagaimana ditunjukkan pada Kode Program 3.4 '
         'berikut.')

caption_par(anchor_close, 'Kode Program 3.4: Aturan Dinamis Jumlah Bullet Point')

code_par(anchor_close,
"""const inputBulletsCount = text
  .split('\\n')
  .map((line) => line.trim())
  .filter((line) => line.length > 0)
  .map((line) => line.replace(/^(•|-|\\*)\\s*/, '').trim())
  .filter((line) => line.length > 0).length;

if (inputBulletsCount > 3) {
  bulletRule = `- You MUST output EXACTLY ${inputBulletsCount} bullet
    points to match the user's input. Do not reduce or combine them.`;
} else {
  bulletRule = `- You MUST output EXACTLY 3 bullet points — no more, no
    less. Expand the user's details into 3 separate bullets based only
    on what the user wrote.`;
}""")

body_par(anchor_close,
         'Berdasarkan Kode Program 3.4, apabila pengguna menuliskan lebih dari tiga poin, '
         'model diwajibkan menghasilkan jumlah poin yang sama persis agar tidak ada '
         'informasi pengguna yang hilang akibat penggabungan poin. Sebaliknya, apabila '
         'pengguna menuliskan tiga poin atau kurang, model diminta mengembangkan masukan '
         'menjadi tepat tiga poin dengan tetap berlandaskan hanya pada informasi yang '
         'ditulis pengguna. Dengan mekanisme ini, panjang deskripsi pada CV tetap terjaga '
         'proporsional tanpa membuka celah bagi model untuk menambahkan informasi baru.')

# ---------------- 3.6.3 ----------------
heading_par(anchor_close, '3.6.3  Prompt Peningkatan Penulisan Mode Professional English')

body_par(anchor_close,
         'Mode *Professional English* ditujukan bagi pengguna yang ingin membuat CV '
         'berbahasa Inggris untuk lamaran kerja internasional. Pada mode ini model tidak '
         'sekadar menerjemahkan, tetapi menerjemahkan sekaligus menulis ulang teks pengguna '
         'menjadi kalimat bahasa Inggris profesional. *System prompt* untuk peningkatan '
         'ringkasan diri pada mode ini ditunjukkan pada Kode Program 3.5 berikut.')

caption_par(anchor_close, 'Kode Program 3.5: System Prompt Peningkatan Ringkasan Diri (Professional English)')

code_par(anchor_close,
"""You are a seasoned career coach who writes like a real person, not a
machine. Translate and rewrite the user's summary into polished
Professional English for an international CV, making it sound like
something a confident human professional would actually write about
themselves.

Rules:
- Write 4-6 well-constructed sentences that flow naturally. Do NOT make
  it too short or overly compressed.
- Use active voice and concrete language. Stay faithful to the original
  meaning — do not add claims or achievements that weren't in the
  source.
- Vary your sentence openings: do NOT start every sentence the same way.
  Mix simple and compound sentences.
- Avoid cliché AI phrases: "proven track record", "passionate about",
  "leveraging", "spearheading", "cutting-edge", "innovative solutions",
  "results-driven", "dynamic professional". Write the way a real person
  talks about their career.
- Keep the tone confident but grounded — not boastful, not robotic.
- Make it ATS-friendly by naturally weaving in relevant keywords from
  the user's input.
- Return ONLY the final summary paragraph. No headings, no labels, no
  conversational filler.""")

body_par(anchor_close,
         'Struktur prompt pada Kode Program 3.5 sama dengan versi Bahasa Indonesia, dengan '
         'dua penyesuaian utama. Pertama, instruksi inti diubah menjadi "*translate and '
         'rewrite*" disertai penekanan "*stay faithful to the original meaning*" agar hasil '
         'terjemahan tidak menyimpang dari makna asli tulisan pengguna. Kedua, daftar '
         'larangan kata klise disesuaikan dengan frasa yang umum muncul pada CV berbahasa '
         'Inggris hasil keluaran AI, seperti "*proven track record*", "*results-driven*", '
         'dan "*dynamic professional*". Untuk peningkatan deskripsi pengalaman pada mode '
         '*English*, digunakan prompt dengan pola yang sama seperti Kode Program 3.3, '
         'termasuk aturan anti-fabrikasi angka dan aturan dinamis jumlah poin pada '
         'Kode Program 3.4, namun dengan keluaran berbahasa Inggris profesional.')

# ---------------- 3.6.4 ----------------
heading_par(anchor_close, '3.6.4  Prompt Impor CV')

body_par(anchor_close,
         'Fitur impor CV memanfaatkan LLM sebagai *CV parser* yang mengubah teks hasil '
         'ekstraksi dokumen PDF/DOCX menjadi data terstruktur dalam format JSON, sehingga '
         'seluruh *form* dapat terisi secara otomatis. Karena keluaran fitur ini harus '
         'dapat dibaca langsung oleh program (*machine-readable*), prompt dirancang jauh '
         'lebih ketat dibandingkan fitur peningkatan penulisan. *System prompt* yang '
         'digunakan ditunjukkan secara ringkas pada Kode Program 3.6 berikut.')

caption_par(anchor_close, 'Kode Program 3.6: System Prompt Impor CV (disajikan secara ringkas)')

code_par(anchor_close,
"""You are an expert CV/Resume parser. Extract all information from the
provided CV text and return it as a valid JSON object.

The JSON must follow this exact schema:
{
  "personalInfo":  { "fullName", "address", "email", "phone",
                     "linkedin", "summary" },
  "educations":    [{ "school", "major", "gpa", "startDate", "endDate",
                     "isCurrent", "description" }],
  "experiences":   [{ "role", "company", "startDate", "endDate",
                     "isCurrent", "description" }],
  "projects":      [{ "name", "role", "startDate", "endDate",
                     "description" }],
  "achievements":  [{ "name", "year" }],
  "skills":        { "hard", "soft" },
  "customSections":[{ "title", "mode", "content", "items": [ ... ] }],
  "sectionOrder":  [ "string" ]
}

Rules:
- Extract ALL information present in the CV
- For descriptions: if the original CV uses bullet points (•, -, *,
  numbers, or any list markers), format EACH bullet as a separate line
  starting with "• ". If the original is a paragraph with no bullets,
  keep it as a plain paragraph.
- For skills: separate hard skills (technical) from soft skills
  (interpersonal) as comma-separated strings
- For dates: keep as-is from the CV (e.g. "Jan 2022", "2020")
- If a field is not found, use empty string "" or empty array []
- Set isCurrent to true if the position says "Present", "Current",
  "Now", "Sekarang", etc.
- IMPORTANT — Standard section mapping: "summary/profile/ringkasan" →
  personalInfo.summary; "education/pendidikan" → educations;
  "experience/pengalaman kerja" → experiences; "projects/proyek" →
  projects; "skills/keahlian" → skills; "awards/prestasi" →
  achievements. These MUST NEVER be placed in "customSections".
- For "customSections": include ONLY sections NOT covered by the
  standard mapping (e.g. Certifications, Languages, Volunteer,
  Organizations, References, Interests, etc.)
- For "sectionOrder": list the section keys in the EXACT ORDER they
  appear in the CV from top to bottom
- Return ONLY the JSON object, no markdown, no explanation, no code
  fences""")

body_par(anchor_close,
         'Berdasarkan Kode Program 3.6, prompt impor CV memiliki empat karakteristik utama. '
         'Pertama, skema JSON didefinisikan secara eksplisit dan lengkap di dalam prompt '
         'sehingga model tidak memiliki kebebasan menentukan struktur keluarannya sendiri. '
         'Kedua, terdapat aturan pemetaan bagian (*section*) standar yang mewajibkan '
         'model memetakan berbagai variasi nama bagian, baik dalam Bahasa Indonesia maupun '
         'Bahasa Inggris (misalnya "ringkasan", "*profile*", "*about me*"), ke dalam bidang '
         'baku yang telah ditentukan, sedangkan bagian di luar pemetaan standar (misalnya '
         'sertifikasi dan bahasa) ditampung pada *customSections* agar tidak ada informasi '
         'yang hilang. Ketiga, aturan *sectionOrder* meminta model mencatat urutan bagian '
         'persis seperti pada dokumen asli sehingga susunan CV pengguna tetap '
         'dipertahankan. Keempat, model diwajibkan mengembalikan hanya objek JSON tanpa '
         '*markdown* maupun penjelasan tambahan agar keluaran dapat langsung diproses '
         'dengan JSON.parse() pada *backend*. Prompt ini dijalankan dengan *temperature* '
         '0,1 sebagaimana dijelaskan pada subbab 3.5 agar keluaran konsisten pada setiap '
         'permintaan.')

# ---------------- 3.6.5 ----------------
heading_par(anchor_close, '3.6.5  Pasca-pemrosesan Keluaran LLM')

body_par(anchor_close,
         'Meskipun format keluaran telah diatur secara ketat di dalam prompt, LLM tidak '
         'selalu mematuhi instruksi secara sempurna. Oleh karena itu, sistem menerapkan '
         'lapisan pasca-pemrosesan (*post-processing*) pada *backend* sebagai pengaman '
         'tambahan sebelum hasil ditampilkan kepada pengguna, sebagaimana ditunjukkan pada '
         'Kode Program 3.7 berikut.')

caption_par(anchor_close, 'Kode Program 3.7: Pasca-pemrosesan Keluaran LLM')

code_par(anchor_close,
"""// Fitur peningkatan penulisan: membersihkan keluaran bullet
return rawResult
  .replace(/\\\\n/g, '\\n')      // ubah literal '\\n' menjadi baris baru
  .split('\\n')
  .map((line) => line.trim())
  .filter((line) => line.length > 0)
  // hapus penanda bullet (•, -, *, penomoran) yang mungkin ditambahkan
  .map((line) => line.replace(/^(?:[•\\-*]|\\d+[.)\\s])\\s*/, '').trim())
  .filter((line) => line.length > 0)
  .join('\\n');

// Fitur impor CV: menghapus code fence markdown pada keluaran JSON
return content
  .replace(/^```(?:json)?\\s*/i, '')
  .replace(/\\s*```$/i, '')
  .trim();""")

body_par(anchor_close,
         'Berdasarkan Kode Program 3.7, pada fitur peningkatan penulisan sistem menghapus '
         'baris kosong, penanda *bullet*, dan penomoran yang terkadang tetap ditambahkan '
         'oleh model, serta mengubah karakter literal "\\n" menjadi baris baru yang '
         'sebenarnya. Pada fitur impor CV, sistem menghapus pagar kode (*code fence*) '
         '*markdown* yang kadang menyelimuti keluaran JSON, kemudian memvalidasi hasilnya '
         'dengan JSON.parse(); apabila validasi gagal, sistem mengembalikan pesan '
         'kesalahan kepada pengguna alih-alih menampilkan data yang rusak. Kombinasi '
         'antara aturan ketat pada prompt dan pasca-pemrosesan ini memastikan keluaran '
         'LLM selalu berada dalam format yang dapat digunakan oleh sistem.')

# ================================================================
# 3) Append 3.7 Metode Pengujian at end of document
# ================================================================

def append_heading(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = SPACE_AFTER
    r = p.add_run(text)
    r.bold = True
    return p


def append_body(text, indent=True):
    p = doc.add_paragraph()
    _fmt_body(p, indent=indent)
    add_runs_with_italic(p, text)
    return p

append_heading('3.7  Metode Pengujian')

append_body('Sesuai dengan tahap *evaluation* pada metode *Design Science Research* '
            '(Tabel 3.1), sistem yang telah dikembangkan dievaluasi melalui serangkaian '
            'pengujian untuk memastikan seluruh fungsi berjalan sesuai kebutuhan dan '
            'keluaran yang dihasilkan berkualitas. Adapun metode pengujian yang digunakan '
            'pada penelitian ini adalah sebagai berikut:')

append_body('1. Pengujian Fungsional (*Black-box Testing*)', indent=False)
append_body('Pengujian *black-box* dilakukan untuk memastikan seluruh fitur sistem, mulai '
            'dari pengisian *form*, impor CV, peningkatan penulisan, pratinjau, hingga '
            'unduh PDF, berfungsi sesuai dengan kebutuhan fungsional yang telah ditetapkan '
            'pada subbab 3.3.1. Pengujian dilakukan dengan menyusun skenario uji (*test '
            'case*) beserta hasil yang diharapkan, kemudian membandingkannya dengan hasil '
            'aktual sistem tanpa memperhatikan struktur internal kode program.')

append_body('2. Pengujian Kualitas Keluaran LLM', indent=False)
append_body('Pengujian ini dilakukan untuk mengevaluasi kualitas hasil fitur berbasis LLM. '
            'Pada fitur peningkatan penulisan, teks masukan pengguna dibandingkan dengan '
            'teks hasil keluaran model untuk menilai perbaikan struktur dan gaya bahasa, '
            'kepatuhan terhadap jumlah poin yang ditentukan, serta kepatuhan terhadap '
            'aturan anti-halusinasi, yaitu tidak munculnya angka maupun pencapaian yang '
            'tidak ditulis pengguna. Pada fitur impor CV, pengujian dilakukan dengan '
            'mengunggah berbagai dokumen CV berformat PDF dan DOCX, kemudian memeriksa '
            'validitas struktur JSON yang dihasilkan serta ketepatan pemetaan informasi '
            'ke dalam *form* CV.')

append_body('3. *User Acceptance Test* (UAT)', indent=False)
append_body('UAT dilakukan dengan melibatkan calon pengguna, yaitu mahasiswa dan pencari '
            'kerja, untuk menggunakan sistem secara langsung dalam menyusun CV. Responden '
            'kemudian mengisi kuesioner penilaian terhadap aspek kemudahan penggunaan, '
            'kebermanfaatan fitur, dan kualitas hasil CV yang dihasilkan. Hasil kuesioner '
            'diolah untuk mengukur tingkat penerimaan pengguna terhadap sistem.')

append_body('4. Pengujian Kompatibilitas', indent=False)
append_body('Pengujian kompatibilitas dilakukan dengan mengakses sistem melalui beberapa '
            'peramban (*browser*) dan perangkat yang berbeda, baik komputer maupun '
            '*smartphone*, untuk memastikan tampilan dan fungsi sistem tetap berjalan '
            'dengan baik sesuai kebutuhan non-fungsional pada subbab 3.3.2.')

append_body('Hasil dari seluruh pengujian tersebut disajikan dan dibahas secara rinci '
            'pada Bab 4.')

doc.save(DST)
print('Saved:', DST)
